"""Generate SPARC System Explanation Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUT = r"d:\ITLS\SPARC\coding-block-SPARC\platform\docs\SPARC_System_Explanation.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)
        r2.italic = True
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_defaults(doc)

    add_title(
        doc,
        "SPARC Current System",
        "Game + Dashboard Explanation\n"
        f"Prepared: {date.today().strftime('%B %d, %Y')}",
    )

    add_para(
        doc,
        "This document describes how the SPARC computational thinking assessment "
        "system works today—from student play in Unity through teacher dashboards—in "
        "plain language suitable for a dissertation proposal, advisor meeting, or job talk.",
    )

    # 1
    add_h1(doc, "1. What SPARC Is")
    add_para(
        doc,
        "SPARC is a robot programming game (Unity WebGL) connected to a Next.js "
        "assessment platform (PostgreSQL). Students build programs with action blocks "
        "(forward, turn, etc.) and press RUN to move Robo on a grid or number line. "
        "The platform records what they did and produces teacher-facing diagnosis—not "
        "just a pass/fail grade.",
    )
    add_table(
        doc,
        ["Part", "Role"],
        [
            ["Unity game", "Play, interaction, animation, local retry logic"],
            ["Platform", "Students/teachers, level catalog, assignments, scoring, reports"],
        ],
    )

    # 2
    add_h1(doc, "2. Big-Picture Architecture")
    add_para(doc, "The system has three layers:")
    add_numbered(
        doc,
        [
            "Student side — login, gamified home, Unity WebGL game",
            "Platform — game APIs, teacher dashboards, assessment engine",
            "Database — levels (items), attempts, events, stealth assessment results",
        ],
    )
    add_para(doc, "Critical design split:", bold=True)
    add_bullets(
        doc,
        [
            "Unity runs the game and sends evidence (commands, flag, touches, time).",
            "Platform re-simulates and scores deterministically on level end—teachers do not depend on Unity physics alone.",
        ],
    )

    # 3
    add_h1(doc, "3. Student Experience")
    add_h2(doc, "3.1 Two Student Interfaces")
    add_table(
        doc,
        ["Path", "Sign-in", "What they use"],
        [
            ["Game path (main)", "Student ID → cookie session", "/student/home → /play (Unity)"],
            [
                "Portal path (optional)",
                "Email/password (NextAuth)",
                "/student/dashboard, progress tables, history",
            ],
        ],
    )
    add_para(
        doc,
        "Most young students use the game path only. The portal is a read-only analytics view.",
    )

    add_h2(doc, "3.2 Typical Play Flow")
    add_numbered(
        doc,
        [
            "Land on /student → enter Student ID (e.g. 1001 → stored as STU-1001).",
            "Home (/student/home) — gamified view: completion %, streak, stars, assigned items, “play next.”",
            "Play (/play) — Unity loads inside the browser (WebGL from platform session).",
            "Unity fetches assigned levels from platform, resumes where they left off.",
            "Student completes items; home stats update after each attempt.",
        ],
    )

    add_h2(doc, "3.3 In-Game Actions (Per Item)")
    add_numbered(
        doc,
        [
            "See the grid or number line, Robo, goal/objects.",
            "Build a program — drag action blocks onto the yellow strip (or edit starter program, fill blanks, place flag).",
            "Press RUN — Robo moves (or answer is checked without animation if configured).",
            "Pass → success message → next item. Fail → retry until maxAttempts.",
            "Reset (if enabled) — clears program and robot position without using an attempt.",
        ],
    )
    add_para(doc, "Intro item (Item 0): step-by-step tutorial for action blocks before free play.")

    # 4
    add_h1(doc, "4. Unity ↔ Platform Connection")
    add_h2(doc, "4.1 Communication")
    add_bullets(
        doc,
        [
            "REST API on the platform (/api/game/*)",
            "Header: X-Game-Api-Key (shared secret)",
            "Student identified by studentId in each payload",
        ],
    )
    add_para(doc, "Key Unity scripts:", bold=True)
    add_bullets(
        doc,
        [
            "PlatformCommunication — login / student create",
            "PlatformLevelLoader — load level list + config JSON",
            "GameAssessmentClient — attempts, events, level end",
            "CharacterMove — gameplay, win/lose, telemetry",
        ],
    )

    add_h2(doc, "4.2 Level Loading")
    add_para(doc, "GET /api/game/levels?studentId=STU-1001 returns:")
    add_bullets(
        doc,
        [
            "Assigned items (ordered)",
            "Full config per item (grid, objects, intro, hints, number line, etc.)",
            "Resume hint (resumeLevelKey, resumeSlot)",
        ],
    )
    add_para(doc, "Assignment rules:", bold=True)
    add_bullets(
        doc,
        [
            "If teacher assigned specific items → only those appear.",
            "Otherwise → teacher’s published catalog.",
            "Hidden/unpublished items are excluded.",
        ],
    )

    # 5
    add_h1(doc, "5. Item Types (Level Types)")
    add_table(
        doc,
        ["Type", "Student task", "Example"],
        [
            ["INTRO", "Guided tutorial for blocks", "Step-by-step forward/turn"],
            ["DRAG_ACTIONS", "Build full program from scratch", "Route Robo to goal / visit objects"],
            ["FLAG_PLACEMENT", "Place flag where Robo will stop", "Prediction before RUN"],
            ["CHOOSE_BUTTONS", "Fill blanks in a fixed program", "Pick correct turn at step 3"],
            ["DRAG_EDIT_PROGRAM", "Fix a broken starter program", "Debug wrong path"],
        ],
    )
    add_para(doc, "Layouts:", bold=True)
    add_bullets(
        doc,
        [
            "GRID — 2D maze-style navigation (default).",
            "NUMBER_LINE — horizontal ticks; counting / step tasks.",
        ],
    )
    add_para(
        doc,
        "The platform picks a different analyzer per type—prediction is not scored like path building.",
    )

    # 6
    add_h1(doc, "6. What Gets Recorded (Telemetry)")
    add_para(doc, "Every RUN creates one attempt on the platform.")
    add_h2(doc, "6.1 During Play")
    add_table(
        doc,
        ["Event", "When", "Why it matters"],
        [
            ["level-start", "Student presses RUN", "Opens attempt, stores initial program"],
            ["save-command-event", "On submit (and optionally each step)", "Full command history"],
            ["save-robot-touch-event", "Student drags/taps Robo", "Spatial exploration evidence"],
            ["save-reset-event", "Reset button", "Persistence / strategy"],
            ["save-progress", "Mid-level save", "Snapshot (secondary)"],
        ],
    )

    add_h2(doc, "6.2 At Level End")
    add_para(doc, "POST /api/game/level-end sends:")
    add_bullets(
        doc,
        [
            "passed, score, status (CORRECT / INCORRECT)",
            "finalCommand — full program string",
            "totalTimeSeconds",
            "mistakes JSON — flag cell, blank answers, visit pattern, etc.",
            "robotTouchCount, resetCount",
            "objectVisit — reached start/end, visitPattern (both, start_only, end_only, neither)",
        ],
    )
    add_para(doc, "One RUN = one attempt. Retries usually create new attempts after reporting the failed run.")

    # 7
    add_h1(doc, "7. Assessment Pipeline (Platform-Side)")
    add_para(
        doc,
        "This is the scientific core—the start of an evidence-centered design (ECD) pipeline.",
    )
    add_numbered(
        doc,
        [
            "Level end payload received from Unity",
            "resolveAttemptEndScore — server re-checks pass/fail",
            "Task type detector — routes to correct analyzer",
            "Simulation / analysis — path, flag, blanks, number line",
            "Evidence extraction — behaviors, mistake types",
            "Teacher narrative — verdict, interpretation, recommendation",
            "Stored in StealthAssessmentResult and shown on dashboard",
        ],
    )

    add_h2(doc, "7.1 Server Re-Scoring by Item Type")
    add_table(
        doc,
        ["Item type", "Server checks"],
        [
            ["DRAG_ACTIONS / DRAG_EDIT", "Simulate program — does Robo stop on goal?"],
            ["FLAG_PLACEMENT", "Does flag match simulated end cell?"],
            ["CHOOSE_BUTTONS", "Are blank answers correct?"],
            ["INTRO", "Completion / guided steps"],
        ],
    )

    add_h2(doc, "7.2 Task-Specific Analyzers")
    add_table(
        doc,
        ["Analyzer", "Detects (examples)"],
        [
            ["Path building", "missingForward, extraForward, wrongRotation, goalOrderError, obstacleCollision"],
            ["Prediction (flag)", "oneStepCountingError, turnAsMoveError, leftRightTurnConfusion"],
            ["Debugging", "Repair quality — over-fix, under-fix, wrong turn"],
            ["Number line", "Step counting, direction, visit order on ticks"],
            ["Choice / blanks", "Per-blank correctness"],
        ],
    )

    add_h2(doc, "7.3 Stealth Assessment Output")
    add_bullets(
        doc,
        [
            "Behaviors — what the student did (observable)",
            "Interpretations — what that might mean",
            "Recommendations — next practice suggestion",
            "Route replay data — path visualization for teachers",
        ],
    )

    add_h2(doc, "7.4 What Is Not Fully Built Yet")
    add_bullets(
        doc,
        [
            "Per-construct CT scores are largely disabled—assessment is task-level, not latent-trait inference.",
            "Dashboard summary scores (0–100 per item) can still feel Canvas-like (passed/failed, average %).",
            "There is no Bayesian student model updating beliefs across attempts yet.",
            "CT construct UI was retired; /teacher/ct-constructs redirects away.",
        ],
    )
    add_para(
        doc,
        "Summary: SPARC has a strong evidence extraction + diagnosis layer. The formal measurement model "
        "(student model ↔ evidence model ↔ probabilistic inference) remains the primary research gap for CaST 2.0.",
        bold=True,
    )

    # 8
    add_h1(doc, "8. Teacher Dashboard (Learning Intelligence)")
    add_para(doc, "Teachers sign in with email/password. They see only their classes, students, and items (scoped access).")
    add_table(
        doc,
        ["Page", "Purpose"],
        [
            ["Dashboard", "Today’s snapshot: active students, attempts, hardest item, who needs support; charts"],
            ["Analytics", "Deeper charts + Excel export"],
            ["Reports", "Class summaries, bulk assessment export"],
            ["Students", "Searchable roster; pass/fail; assign items; open profiles"],
            ["Classes", "Roster + class-wide item assignment"],
            ["Items", "Create/edit/publish levels; view per-item pass rates"],
            ["Attempt detail", "Full evidence for one RUN"],
        ],
    )

    add_h2(doc, "8.1 Per-Attempt Detail (Deepest View)")
    add_para(doc, "/teacher/attempts/[id] shows:")
    add_numbered(
        doc,
        [
            "Verdict banner — plain-English headline (e.g. “Robo stopped one cell short…”).",
            "Task-specific panel — path map, flag prediction, blank answers, number line.",
            "Command timeline — what they built and when.",
            "Route replay — student path vs reference path (where applicable).",
            "Mistake type — classified error (e.g. goalOrderError).",
            "Stealth assessment — behaviors, interpretation, recommendation.",
            "Teacher notes — editable comments.",
        ],
    )
    add_para(
        doc,
        "Teacher value today: “What happened and what to do next”—grounded in simulation, not only a percentage score.",
    )

    # 9
    add_h1(doc, "9. Student Dashboard (Portal — Secondary)")
    add_table(
        doc,
        ["Page", "Shows"],
        [
            ["Dashboard", "Summary cards — completion, passed/failed/incomplete"],
            ["Progress", "Table of all items, scores, time"],
            ["Items", "Card per item with status"],
            ["Item detail", "Best attempt, command, feedback text"],
            ["History", "Chronological attempt list"],
        ],
    )
    add_para(doc, "Students do not see route replay or mistake taxonomy—that is teacher-facing.")

    # 10
    add_h1(doc, "10. Teacher Workflow (End-to-End)")
    add_numbered(
        doc,
        [
            "Teacher builds items (level editor) or uses seed catalog.",
            "Assigns items to class or individual students.",
            "Students play in Unity; data streams automatically.",
            "Teacher monitors dashboard (who is struggling).",
            "Opens attempt detail for diagnosis.",
            "Exports Excel for reporting / research.",
        ],
    )

    # 11
    add_h1(doc, "11. Relation to Dissertation / CaST 2.0")
    add_table(
        doc,
        ["Advisor concern", "Current system"],
        [
            ["Canvas-like points", "Item scores and pass rates exist; attempt detail is richer than Canvas"],
            ["Measurement model", "Evidence model is partial — behaviors + mistake types per task"],
            ["Student model", "Not formalized — no latent traits updated across items"],
            ["Digital replication of CaST", "Yes — game items + logging + teacher diagnosis"],
            ["CaST 2.0 / digital-native", "Partially — simulation, replay, process logs; missing probabilistic inference"],
        ],
    )
    add_para(
        doc,
        "Proposal sentence: SPARC’s current platform implements a task-aware evidence pipeline "
        "(gameplay → simulation → mistake classification → teacher narrative), which supports "
        "classroom use and CaST comparison, but the measurement model—explicit latent skills and "
        "probabilistic inference from observables—remains the primary research contribution for CaST 2.0.",
        bold=True,
    )

    # 12
    add_h1(doc, "12. System Strengths")
    add_bullets(
        doc,
        [
            "Digital affordances: full command history, route replay, reset/touch counts, time on task.",
            "Deterministic scoring: same program → same diagnosis (reproducible for research).",
            "Task-aware assessment: flag ≠ path ≠ blanks ≠ number line.",
            "Teacher-usable output: plain-language verdicts, not raw logs.",
            "Assignment + scope: real classroom deployment model.",
            "Authoring: teachers can build items without Unity rebuild.",
        ],
    )

    # 13
    add_h1(doc, "13. Known Gaps")
    add_bullets(
        doc,
        [
            "No unified student model (latent CT/math skills) across items.",
            "Item scores aggregate behavior into single numbers without uncertainty.",
            "Construct weighting removed — no Q-matrix linking items to skills.",
            "Some telemetry defined but unused (e.g. action button events).",
            "Math units (counting, geometry) not yet a separate catalog—but grid/number-line mechanics support them.",
        ],
    )

    # 14
    add_h1(doc, "14. Glossary")
    add_table(
        doc,
        ["Term", "Meaning in SPARC"],
        [
            ["Item", "One level/unit of play (database: Level)"],
            ["Attempt", "One RUN on an item (LevelAttempt)"],
            ["Slot", "Position in student’s assigned list (1, 2, 3…)"],
            ["Stealth assessment", "Automatic diagnosis written after each attempt"],
            ["Evidence", "Commands, path, flag, visits, edits, time"],
            ["Verdict", "Top-line teacher summary on attempt page"],
        ],
    )

    # footer
    doc.add_page_break()
    add_h2(doc, "Document Information")
    add_para(doc, "Project: SPARC Computational Thinking Assessment Platform")
    add_para(doc, "Repository: coding-block-SPARC (Unity + platform/)")
    add_para(doc, f"Generated: {date.today().isoformat()}")
    add_para(
        doc,
        "For questions about assessment architecture, see also: "
        "CT_Assessment_System_Architecture.md and platform/docs/STEALTH_ASSESSMENT.md",
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
